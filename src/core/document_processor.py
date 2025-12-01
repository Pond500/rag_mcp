"""Document Processor

Extracts text from various file formats (PDF, DOCX, TXT) and chunks them.
"""
from __future__ import annotations
from typing import List, Dict, Any, Optional
from pathlib import Path
import logging

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process documents: extract text and chunk
    
    Usage:
        processor = DocumentProcessor(config)
        pages = processor.extract_text(file_path="doc.pdf")
        chunks = processor.chunk_text(pages, chunk_size=1000, overlap=200)
    """
    
    def __init__(self, config=None):
        self.config = config
        self.chunk_size = getattr(config, "chunk_size", 1000) if config else 1000
        self.chunk_overlap = getattr(config, "chunk_overlap", 200) if config else 200
    
    def extract_text(self, file_path: str, file_content: Optional[bytes] = None) -> List[str]:
        """Extract text from file (supports PDF, DOCX, TXT)
        
        Args:
            file_path: Path to file (used for extension detection)
            file_content: Optional raw bytes (if provided, will use instead of reading file)
            
        Returns:
            List of text strings (one per page for PDF/DOCX, or single string for TXT)
        """
        ext = Path(file_path).suffix.lower()
        
        try:
            if ext == ".pdf":
                return self._extract_pdf(file_path, file_content)
            elif ext == ".docx":
                return self._extract_docx(file_path, file_content)
            elif ext in [".txt", ".md"]:
                return self._extract_text(file_path, file_content)
            else:
                logger.warning("Unsupported file type: %s, treating as text", ext)
                return self._extract_text(file_path, file_content)
        except Exception as e:
            logger.error("Failed to extract text from %s: %s", file_path, e)
            return []
    
    def _extract_pdf(self, file_path: str, file_content: Optional[bytes]) -> List[str]:
        """Extract text from PDF"""
        try:
            import fitz  # PyMuPDF
            
            if file_content:
                doc = fitz.open(stream=file_content, filetype="pdf")
            else:
                doc = fitz.open(file_path)
            
            pages = []
            for page in doc:
                text = page.get_text()
                if text.strip():
                    pages.append(text)
            
            doc.close()
            logger.info("Extracted %d pages from PDF", len(pages))
            return pages
            
        except ImportError:
            logger.warning("PyMuPDF not available, falling back to simple extraction")
            # Fallback: return dummy text
            return ["(PDF content - PyMuPDF not installed)"]
        except Exception as e:
            logger.error("PDF extraction failed: %s", e)
            return []
    
    def _extract_docx(self, file_path: str, file_content: Optional[bytes]) -> List[str]:
        """Extract text from DOCX"""
        try:
            from docx import Document
            from io import BytesIO
            
            if file_content:
                doc = Document(BytesIO(file_content))
            else:
                doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Combine into single page
            full_text = "\n".join(paragraphs)
            logger.info("Extracted DOCX with %d paragraphs", len(paragraphs))
            return [full_text] if full_text else []
            
        except ImportError:
            logger.warning("python-docx not available")
            return ["(DOCX content - python-docx not installed)"]
        except Exception as e:
            logger.error("DOCX extraction failed: %s", e)
            return []
    
    def _extract_text(self, file_path: str, file_content: Optional[bytes]) -> List[str]:
        """Extract text from TXT/MD"""
        try:
            if file_content:
                text = file_content.decode('utf-8', errors='ignore')
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            
            return [text] if text.strip() else []
            
        except Exception as e:
            logger.error("Text extraction failed: %s", e)
            return []
    
    def chunk_text(
        self,
        pages: List[str],
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None
    ) -> List[Dict[str, Any]]:
        """Chunk text with overlap
        
        Args:
            pages: List of page texts
            chunk_size: Characters per chunk (uses config default if None)
            chunk_overlap: Overlap between chunks (uses config default if None)
            
        Returns:
            List of chunks with metadata: [{"text": "...", "page": 1, "chunk_index": 0}, ...]
        """
        chunk_size = chunk_size or self.chunk_size
        chunk_overlap = chunk_overlap or self.chunk_overlap
        
        chunks = []
        chunk_index = 0
        
        for page_num, page_text in enumerate(pages, start=1):
            # Simple character-based chunking with overlap
            start = 0
            while start < len(page_text):
                end = start + chunk_size
                chunk_text = page_text[start:end]
                
                if chunk_text.strip():
                    chunks.append({
                        "text": chunk_text,
                        "page": page_num,
                        "chunk_index": chunk_index,
                        "char_start": start,
                        "char_end": min(end, len(page_text))
                    })
                    chunk_index += 1
                
                # Move forward (with overlap)
                start += chunk_size - chunk_overlap
                
                # Break if we're at the end
                if end >= len(page_text):
                    break
        
        logger.info("Created %d chunks from %d pages", len(chunks), len(pages))
        return chunks
